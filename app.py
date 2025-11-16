import os
from io import BytesIO
from datetime import datetime
from functools import wraps

from flask import Flask, render_template, redirect, url_for, flash, request, send_file, abort
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from itsdangerous import URLSafeTimedSerializer

import smtplib
from email.mime.text import MIMEText

from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# Load env from .env **only locally**
if os.getenv("VERCEL") is None:
    from dotenv import load_dotenv
    load_dotenv()

# Flask app
app = Flask(__name__)
app.config["SECRET_KEY"] = os.getenv("SECRET_KEY")

# Database config
app.config["SQLALCHEMY_DATABASE_URI"] = (
    f"mysql+pymysql://{os.getenv('DB_USER')}:"
    f"{os.getenv('DB_PASSWORD')}@"
    f"{os.getenv('DB_HOST')}:{os.getenv('DB_PORT')}/"
    f"{os.getenv('DB_NAME')}"
)
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

# Initialize DB
from models import db, User, Course, Enrollment, Blog, CV  # adjust your models import
db.init_app(app)

# Login manager
login_manager = LoginManager()
login_manager.login_view = "login"
login_manager.init_app(app)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Admin decorator
def admin_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not current_user.is_authenticated or not getattr(current_user, "is_admin", False):
            flash(" 🚫 You do not have permission to access this page.", "danger")
            return redirect(url_for("home"))
        return f(*args, **kwargs)
    return wrapper

# Email verification setup
s = URLSafeTimedSerializer(app.config["SECRET_KEY"])

def send_verification_email(user_email, token):
    sender = os.getenv("MAIL_USERNAME")
    sender_pw = os.getenv("MAIL_PASSWORD")
    base = os.getenv("BASE_URL", "http://127.0.0.1:5000")  # fallback to localhost

    link = f"{base}/verify/{token}"
    msg = MIMEText(f"Click to verify your account: {link}")
    msg["Subject"] = "Verify your account"
    msg["From"] = sender
    msg["To"] = user_email

    try:
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(sender, sender_pw)
            server.sendmail(sender, user_email, msg.as_string())
    except Exception as e:
        print("Error sending email:", e)


# Routes

@app.route("/")
def home():
    courses = Course.query.limit(4).all()
    testimonials = [
        {"name": "Aarav", "role": "Student", "feedback": "This platform gave me practical skills!", "avatar": url_for("static", filename="images/user1.jpg")},
        {"name": "Neha", "role": "Data Analyst", "feedback": "Feels like PW Skills but more tailored!", "avatar": url_for("static", filename="images/user2.jpg")},
        {"name": "Rohit", "role": "Developer", "feedback": "Instructors are highly experienced.", "avatar": url_for("static", filename="images/user3.jpg")},
    ]
    return render_template("home.html", courses=courses, testimonials=testimonials, current_year=datetime.now().year)


@app.route("/course/<int:course_id>")
def course_detail(course_id):
    course = Course.query.get_or_404(course_id)
    return render_template("course_detail.html", course=course)


@app.route("/enroll/<int:course_id>")
@login_required
def enroll(course_id):
    course = Course.query.get_or_404(course_id)
    existing = Enrollment.query.filter_by(user_id=current_user.id, course_id=course_id).first()
    if existing:
        flash("Already enrolled ✅", "info")
    else:
        en = Enrollment(user_id=current_user.id, course_id=course_id)
        db.session.add(en)
        db.session.commit()
        flash(f"Enrolled in {course.title} 🎉", "success")
    return redirect(url_for("course_detail", course_id=course_id))


@app.route("/my-courses")
@login_required
def my_courses():
    enrolls = Enrollment.query.filter_by(user_id=current_user.id).all()
    courses = [e.course for e in enrolls]
    return render_template("my_courses.html", courses=courses)



@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form.get("username")
        email = request.form.get("email")
        password = request.form.get("password")

        if User.query.filter_by(email=email).first():
            flash("Email already exists!", "danger")
            return redirect(url_for("register"))

        hashed = generate_password_hash(password, method="pbkdf2:sha256", salt_length=8)
        user = User(username=username, email=email, password=hashed, is_verified=False)
        db.session.add(user)
        db.session.commit()

        token = s.dumps(email, salt="email-confirm")
        send_verification_email(email, token)
        flash("Account created! Check email for verification.", "info")
        return redirect(url_for("login"))

    return render_template("register.html")


@app.route("/verify/<token>")
def verify_email(token):
    try:
        email = s.loads(token, salt="email-confirm", max_age=3600)
        user = User.query.filter_by(email=email).first()
        if user:
            user.is_verified = True
            db.session.commit()
            flash("✅ Email verified! Log in now.", "success")
            return redirect(url_for("login"))
    except Exception as e:
        print("Verify error:", e)
        flash("Invalid or expired link.", "danger")
    return redirect(url_for("home"))


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email")
        pw = request.form.get("password")
        user = User.query.filter_by(email=email).first()
        if user and check_password_hash(user.password, pw):
            if not user.is_verified:
                flash("Verify your email first.", "warning")
                return redirect(url_for("login"))
            login_user(user)
            flash(f"Welcome {user.username}!", "success")
            return redirect(url_for("dashboard"))
        flash("Invalid email/password.", "danger")
    return render_template("login.html")


@app.route("/logout")
@login_required
def logout():
    logout_user()
    flash("Logged out.", "info")
    return redirect(url_for("home"))


@app.route("/dashboard")
@login_required
def dashboard():
    return render_template("dashboard.html", user=current_user)


@app.route("/admin")
@login_required
@admin_required
def admin_panel():
    users = User.query.all()
    return render_template("admin.html", users=users)


@app.route("/admin/update_role/<int:user_id>", methods=["POST"])
@login_required
@admin_required
def update_role(user_id):
    user = User.query.get(user_id)
    new_role = request.form.get("role")
    if user and new_role in ["student", "instructor", "admin"]:
        user.role = new_role
        db.session.commit()
        flash("Role updated!", "success")
    return redirect(url_for("admin_panel"))


@app.route("/courses")
def courses():
    search = request.args.get("search", "")
    category = request.args.get("category", "")
    sort = request.args.get("sort", "")

    q = Course.query
    if search:
        q = q.filter(Course.title.ilike(f"%{search}%"))
    if category:
        q = q.filter_by(category=category)
    if sort == "price":
        q = q.order_by(Course.price.asc())
    elif sort == "rating":
        q = q.order_by(Course.rating.desc())

    courses = q.all()
    return render_template("courses.html", courses=courses)


@app.route("/update-progress/<int:enrollment_id>/<int:progress>")
@login_required
def update_progress(enrollment_id, progress):
    enrollment = Enrollment.query.get_or_404(enrollment_id)
    if enrollment.user_id != current_user.id:
        abort(403)
    enrollment.progress = min(progress, 100)
    db.session.commit()
    flash("Progress updated!", "success")
    return redirect(url_for("my_courses"))


@app.route("/blogs")
def blogs():
    blogs = Blog.query.order_by(Blog.created_at.desc()).all()
    return render_template("blogs.html", blogs=blogs)


@app.route("/add_blog", methods=["GET", "POST"])
@login_required
def add_blog():
    if request.method == "POST":
        title = request.form.get("title")
        content = request.form.get("content")
        new = Blog(title=title, content=content, author=current_user.username)
        db.session.add(new)
        db.session.commit()
        flash("Blog created!", "success")
        return redirect(url_for("blogs"))
    return render_template("add_blog.html")


@app.route("/blogs/<int:blog_id>/edit", methods=["GET", "POST"])
@login_required
def edit_blog(blog_id):
    blog = Blog.query.get_or_404(blog_id)
    if request.method == "POST":
        blog.title = request.form.get("title")
        blog.content = request.form.get("content")
        db.session.commit()
        flash("Blog updated!", "success")
        return redirect(url_for("blogs"))
    return render_template("edit_blog.html", blog=blog)


@app.route("/blogs/<int:blog_id>/delete", methods=["POST"])
@login_required
def delete_blog(blog_id):
    blog = Blog.query.get_or_404(blog_id)
    db.session.delete(blog)
    db.session.commit()
    flash("Blog deleted!", "danger")
    return redirect(url_for("blogs"))


@app.route("/blogs/<int:blog_id>")
def blog_detail(blog_id):
    blog = Blog.query.get_or_404(blog_id)
    return render_template("blog_detail.html", blog=blog)


@app.route("/cv-builder", methods=["GET", "POST"])
@login_required
def cv_builder():
    if request.method == "POST":
        data = request.form
        cv = CV(
            user_id=current_user.id,
            full_name=data.get("full_name"),
            email=data.get("email"),
            phone=data.get("phone"),
            summary=data.get("summary"),
            skills=data.get("skills"),
            experience=data.get("experience"),
            education=data.get("education"),
            projects=data.get("projects"),
        )
        db.session.add(cv)
        db.session.commit()
        flash("CV Saved!", "success")
        return redirect(url_for("my_cv", cv_id=cv.id))
    return render_template("cv_builder.html")


@app.route("/cv/<int:cv_id>")
@login_required
def my_cv(cv_id):
    cv = CV.query.get_or_404(cv_id)
    return render_template("cv_preview.html", cv=cv)


@app.route("/cv/<int:cv_id>/download/<format>")
@login_required
def download_cv(cv_id, format):
    cv = CV.query.get_or_404(cv_id)

    # PDF
    if format.lower() == "pdf":
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer)
        styles = getSampleStyleSheet()
        elems = []
        elems.append(Paragraph(cv.full_name, styles["Title"]))
        elems.append(Spacer(1, 12))
        elems.append(Paragraph(f"Email: {cv.email}", styles["Normal"]))
        elems.append(Paragraph(f"Phone: {cv.phone}", styles["Normal"]))
        elems.append(Spacer(1, 12))
        for section_name, content in [
            ("Summary", cv.summary),
            ("Skills", cv.skills),
            ("Experience", cv.experience),
            ("Education", cv.education),
            ("Projects", cv.projects),
        ]:
            elems.append(Paragraph(section_name, styles["Heading2"]))
            elems.append(Paragraph(content or "", styles["Normal"]))
            elems.append(Spacer(1, 12))
        doc.build(elems)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name=f"{cv.full_name}.pdf", mimetype="application/pdf")

    # DOCX
    elif format.lower() == "docx":
        buffer = BytesIO()
        docx = Document()
        docx.add_heading(cv.full_name, 0)
        docx.add_paragraph(f"Email: {cv.email}")
        docx.add_paragraph(f"Phone: {cv.phone}")
        for section_name, content in [
            ("Summary", cv.summary),
            ("Skills", cv.skills),
            ("Experience", cv.experience),
            ("Education", cv.education),
            ("Projects", cv.projects),
        ]:
            docx.add_heading(section_name, level=1)
            docx.add_paragraph(content or "")
        docx.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{cv.full_name}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    # TXT
    elif format.lower() == "txt":
        buffer = BytesIO()
        content = f"""{cv.full_name}
Email: {cv.email}
Phone: {cv.phone}

Summary:
{cv.summary or ""}

Skills:
{cv.skills or ""}

Experience:
{cv.experience or ""}

Education:
{cv.education or ""}

Projects:
{cv.projects or ""}
"""
        buffer.write(content.encode("utf-8"))
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name=f"{cv.full_name}.txt", mimetype="text/plain")

    else:
        flash("Invalid format", "danger")
        return redirect(url_for("my_cv", cv_id=cv.id))


# Error handlers
@app.errorhandler(404)
def not_found(e):
    return render_template("404.html"), 404

@app.errorhandler(500)
def server_error(e):
    return render_template("500.html"), 500

@app.errorhandler(403)
def forbidden(e):
    return render_template("403.html"), 403


# No app.run() for Vercel — but allow local run
if __name__ == "__main__":
    with app.app_context():
        db.create_all()
    app.run(debug=True)
